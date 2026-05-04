"""
🎤 발표문(Speech Script) 생성기.

용도:
- 사장님이 발표 리허설 때 보고 거의 암기할 수 있는 시나리오.
- 실제 발표장에서는 PT만 보고 진행하므로, 사전 암기 목적의 도입·본문·마무리 풀멘트 + 예상 Q&A.

산출물:
- DOCX (인쇄 친화 — 큰 폰트, 슬라이드 번호 매핑, 멘트/제스처/시간 가이드 분리)
- 발표시간 15~20분 + Q&A 10~15분 분량 (총 한국어 약 4500~6000자)
"""

from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from loguru import logger

from config.settings import MODELS, OUTPUT_DIR, get_effective_company
from schemas.models import BidNotice, ProposalDraft
from tools.llm_client import chat_anthropic

SYSTEM_PROMPT = """당신은 한국 공공·민간 입찰 PT를 100건 이상 직접 발표한 시니어 사업개발 임원입니다.
박제안 팀의 본문과 최피티의 슬라이드 흐름을 받아, 사장님이 리허설 때 보고 암기할 발표 시나리오를 작성합니다.

원칙:
1. 발표 본 시간은 사용자가 지정한 분량(예: 15~20분). 이를 4개 블록으로 나눈다 — 도입(2~3분) / 사업이해(3~4분) / 솔루션·방법론(7~9분) / 마무리(2~3분).
2. 각 슬라이드 단위로 '몇 분 몇 초', '핵심 멘트(반드시 말할 한 문장)', '풀멘트(읽을 그대로의 문장들 — 자연스러운 구어)', '제스처/포인터/숨고르기 가이드'를 명시한다.
3. 상투적 표현 금지 — '저희가 ~하겠습니다' 보다 '~ 합니다', '~ 가능합니다' 같이 단정형.
4. 회사 자랑 1회 이내, 발주처 입장 공감 3회 이상.
5. 마지막에 '예상 Q&A 10건' 추가 — 각각 질문·핵심답변·근거·5초 이내 마무리.
6. 출력은 마크다운(### / 표/리스트)만 사용. 다른 메타설명 금지."""


def _build_prompt(bid: BidNotice, draft: ProposalDraft, minutes: int) -> str:
    company = get_effective_company()
    sections_text = "\n\n".join(
        f"## {sec.title}\n{(sec.body or '')[:1500]}"
        for sec in sorted(draft.sections, key=lambda s: s.order)
    )
    return f"""다음 입찰 PT의 발표 시나리오를 작성하시오.

# 사업 정보
- 사업명: {bid.title}
- 발주기관: {bid.agency}
- 예산: {bid.budget_krw or '미공개'}원
- 사업기간: {bid.duration_months or '미상'}개월

# 제안사
- 회사명: {company.name}
- 대표: {company.ceo}
- 한 줄 차별화: {company.differentiators or '(미지정)'}

# 발표 분량
- 본 발표: {minutes}분 (도입·사업이해·솔루션방법론·마무리 4블록)
- Q&A: 약 10~15분 (예상 질문 10건)

# 박제안 팀이 작성한 본문 섹션 (요약)
{sections_text}

# 출력 마크다운 구조 (반드시 그대로 따를 것)

# {bid.title} — 발표 시나리오

## 0. 발표자 자기 점검 체크리스트 (리허설 1회차 전)
- [ ] 1번 슬라이드 첫 멘트 암기
- [ ] 시스템·서버 아키텍처 슬라이드에서 카메라/포인터 위치
- [ ] 마무리 한 문장(헤드라인) 암기
- [ ] 예상 Q&A 1·2·3번 답변 암기

## 1. 도입 (00:00 ~ 02:30, 약 2분 30초)

### S01. 표지 — 인사 + 회사 소개 한 줄
- ⏱ 00:00 ~ 00:25 (25초)
- 🎯 핵심 한 문장: "..."
- 🗣 풀멘트:
  > (인사) 안녕하십니까. {company.name} 대표 {company.ceo}입니다. ...
- 👉 제스처: 시선은 가운데 심사위원 → 좌우 1초씩.

(이런 식으로 슬라이드 단위로 모두 작성)

## 2. 사업 이해 (02:30 ~ 06:00)
... (계속) ...

## 3. 솔루션 · 방법론 (06:00 ~ 14:00)
... (계속, 시스템 아키텍처·서버 아키텍처 슬라이드는 멘트가 더 풍부하게) ...

## 4. 마무리 (14:00 ~ {minutes}:00)
... (계속) ...

## 5. 예상 Q&A 10건

### Q1. (질문)
- 💬 핵심 답변 (5초 이내):
- 🧱 근거:
- 🔚 마무리 한 줄:

(Q2 ~ Q10 동일 형식)

## 6. 발표 직전 30초 셀프 체크
- 마이크 / 포인터 / 슬라이드 제어 / 비상용 백업 USB
- 발주처 담당자 이름 1회 호명할 타이밍 (어느 슬라이드?)
"""


def _ask_llm(bid: BidNotice, draft: ProposalDraft, minutes: int) -> str:
    return chat_anthropic(
        MODELS.writer,
        SYSTEM_PROMPT,
        _build_prompt(bid, draft, minutes),
        max_tokens=8000,
    )


# ---------------------------------------------------------------------------
# DOCX 렌더링
# ---------------------------------------------------------------------------

def _add_heading(doc, text: str, level: int = 1, color=RGBColor(0x0B, 0x2C, 0x5A)) -> None:
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(4)
    run = h.add_run(text)
    run.font.name = "맑은 고딕"
    sizes = {1: 22, 2: 18, 3: 15, 4: 13}
    run.font.size = Pt(sizes.get(level, 13))
    run.font.bold = True
    run.font.color.rgb = color


def _add_para(doc, text: str, *, size: int = 13, bold: bool = False,
              color=RGBColor(0x1F, 0x29, 0x37), italic: bool = False,
              indent_cm: float = 0.0) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    run = p.add_run(text)
    run.font.name = "맑은 고딕"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def _markdown_to_docx(doc: Document, md: str) -> None:
    """간단한 마크다운 → DOCX 변환."""
    for raw_line in md.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            doc.add_paragraph()
            continue
        # 헤딩
        m = re.match(r"^(#{1,4})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            _add_heading(doc, m.group(2).strip(), level=level)
            continue
        # 인용(풀멘트)
        if line.lstrip().startswith(">"):
            _add_para(doc, line.lstrip()[1:].strip(),
                      size=14, italic=False,
                      color=RGBColor(0x1E, 0x6F, 0xE6),
                      indent_cm=0.5, bold=True)
            continue
        # 체크박스/리스트
        if re.match(r"^\s*-\s*\[[ xX]\]\s+", line):
            _add_para(doc, line.strip(), size=12, indent_cm=0.5)
            continue
        if line.lstrip().startswith("- "):
            _add_para(doc, "•  " + line.lstrip()[2:].strip(), size=12, indent_cm=0.5)
            continue
        # 시간/이모지로 시작하는 가이드 줄
        _add_para(doc, line.strip(), size=12)


def render_script(bid: BidNotice, draft: ProposalDraft,
                  minutes: int = 18) -> Path:
    """발표문 DOCX 생성. minutes = 본 발표 분량(분)."""
    md = _ask_llm(bid, draft, minutes)
    company = get_effective_company()
    doc = Document()
    # 페이지 여백 2cm
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 표지
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run("🎤 발표 시나리오 (리허설용)")
    r.font.name = "맑은 고딕"
    r.font.size = Pt(26)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x0B, 0x2C, 0x5A)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run(bid.title)
    r2.font.name = "맑은 고딕"
    r2.font.size = Pt(18)
    r2.font.bold = True
    r2.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm = meta.add_run(
        f"{bid.agency}  |  본 발표 {minutes}분 + Q&A 10~15분  |  "
        f"{company.name} · 대표 {company.ceo}  |  "
        f"리허설 작성일 {datetime.now().strftime('%Y-%m-%d')}"
    )
    rm.font.name = "맑은 고딕"
    rm.font.size = Pt(11)
    rm.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

    doc.add_paragraph()  # spacer
    _markdown_to_docx(doc, md)

    fname = f"{bid.bid_id}_speech_v{draft.version}.docx"
    out_path = OUTPUT_DIR / fname
    doc.save(out_path)
    logger.info(f"🎤 발표문 DOCX 생성: {out_path}")
    return out_path
