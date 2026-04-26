"""
DOCX 제안서 생성기.
python-docx로 한국 공공제안서 표준 목차에 맞춰 출력.
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from loguru import logger

from config.settings import OUTPUT_DIR, get_effective_company
from schemas.models import ProposalDraft


def render_proposal_docx(draft: ProposalDraft, title: str) -> Path:
    """제안서 초안을 DOCX 파일로 저장하고 경로를 반환."""
    company = get_effective_company()
    doc = Document()

    # 표지
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("제 안 서")
    run.bold = True
    run.font.size = Pt(36)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(20)
    run.bold = True

    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"제출일자: {datetime.now().strftime('%Y년 %m월 %d일')}").font.size = Pt(14)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"제안사: {company.name}").font.size = Pt(14)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"대표이사: {company.ceo}").font.size = Pt(14)

    doc.add_page_break()

    # 목차
    doc.add_heading("목   차", level=1)
    for i, sec in enumerate(sorted(draft.sections, key=lambda s: s.order), 1):
        doc.add_paragraph(f"{i}. {sec.title}", style="List Number")
    doc.add_page_break()

    # 본문
    for sec in sorted(draft.sections, key=lambda s: s.order):
        doc.add_heading(sec.title, level=1)
        for para in sec.body.split("\n\n"):
            para = para.strip()
            if not para:
                continue
            if para.startswith("- ") or para.startswith("• "):
                for line in para.split("\n"):
                    line = line.strip().lstrip("-•").strip()
                    if line:
                        doc.add_paragraph(line, style="List Bullet")
            else:
                doc.add_paragraph(para)
        doc.add_page_break()

    fname = f"{draft.bid_id}_proposal_v{draft.version}.docx"
    out_path = OUTPUT_DIR / fname
    doc.save(out_path)
    logger.info(f"DOCX 생성: {out_path}")
    return out_path
